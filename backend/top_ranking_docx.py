"""DOCX export helpers for GTI Top Rankings results."""

from __future__ import annotations

import json
import re
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from lxml import etree as _etree
from docx import Document
from docx.opc.packuri import PackURI as _PackURI
from docx.opc.part import Part as _Part
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate, InlineImage


EMPTY_RANKING_NOTE = (
    "Field not present in GTI Intelligence Search preview for this sample."
)

RANKING_LABELS = {
    "top_industries": "Top targeted industries",
    "top_targeted_regions": "Top targeted regions",
    "top_source_regions": "Top source regions",
    "top_tags": "Top tags / themes",
    "threat_categories": "Threat categories",
    "collection_type_distribution": "Collection type distribution",
    "timeline": "Timeline",
    "top_targeted_organizations": "Top targeted organizations",
    "top_tactics": "Top MITRE tactics",
    "top_techniques": "Top MITRE techniques",
    "top_subtechniques": "Top MITRE subtechniques",
}

RANKING_RESULT_KEYS = {
    "top_industries": "targeted_industries",
    "top_targeted_regions": "targeted_regions",
    "top_source_regions": "source_regions",
    "top_tags": "tags",
    "threat_categories": "threat_categories",
    "collection_type_distribution": "collection_type",
    "timeline": "timeline",
    "top_targeted_organizations": "targeted_organizations",
}

CROSS_ANALYSIS_LABELS = {
    "industries_by_tags": "Industries by tags / themes",
    "industries_by_collection_type": "Industries by collection type",
    "industries_by_targeted_region": "Industries by targeted region",
    "timeline_by_collection_type": "Timeline by collection type",
    "source_region_by_targeted_region": "Source region by targeted region",
}


_CHART_CT = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
_CHART_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
_CHART_PALETTE = [
    "0D7F7A", "1A9E98", "0A5F5B", "15B0A8",
    "086560", "20C4BB", "053E3B", "25D4CA",
]
# Semantic severity colors for multi-series charts
_SEVERITY_PALETTE = {
    "high": "DC3545",
    "medium": "FD7E14",
    "low": "28A745",
    "risk": "0D7F7A",
    "noise": "6C757D",
    "success": "28A745",
    "skip": "FFC107",
    "error": "DC3545",
}
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_C = "http://schemas.openxmlformats.org/drawingml/2006/chart"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _xe(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _emu(inches: float) -> int:
    return round(inches * 914_400)


def _build_bar_chart_xml(
    labels: list[str],
    values: list[int],
    title: str,
    horizontal: bool = True,
) -> str:
    n = len(labels)
    bar_dir = "bar" if horizontal else "col"
    cat_pos = "l" if horizontal else "b"
    val_pos = "b" if horizontal else "l"
    orientation = "maxMin" if horizontal else "minMax"
    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(l))}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    pts_val = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
        for i, v in enumerate(values)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:barChart>
        <c:barDir val="{bar_dir}"/>
        <c:grouping val="clustered"/>
        <c:varyColors val="0"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          <c:spPr>
            <a:solidFill><a:srgbClr val="0D7F7A"/></a:solidFill>
            <a:ln><a:solidFill><a:srgbClr val="0A5F5B"/></a:solidFill></a:ln>
          </c:spPr>
          <c:dLbls>
            <c:numFmt formatCode="General" sourceLinked="0"/>
            <c:showLegendKey val="0"/><c:showVal val="1"/>
            <c:showCatName val="0"/><c:showSerName val="0"/>
            <c:showPercent val="0"/><c:showBubbleSize val="0"/>
          </c:dLbls>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
        </c:ser>
        <c:axId val="100"/><c:axId val="101"/>
      </c:barChart>
      <c:catAx>
        <c:axId val="100"/>
        <c:scaling><c:orientation val="{orientation}"/></c:scaling>
        <c:delete val="0"/><c:axPos val="{cat_pos}"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="101"/>
      </c:catAx>
      <c:valAx>
        <c:axId val="101"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="{val_pos}"/>
        <c:numFmt formatCode="General" sourceLinked="0"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="100"/>
        <c:crossBetween val="between"/>
      </c:valAx>
    </c:plotArea>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _build_pie_chart_xml(labels: list[str], values: list[int], title: str) -> str:
    n = len(labels)
    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(l))}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    pts_val = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
        for i, v in enumerate(values)
    )
    color_pts = "".join(
        f'<c:dPt><c:idx val="{i}"/><c:spPr>'
        f'<a:solidFill><a:srgbClr val="{_CHART_PALETTE[i % len(_CHART_PALETTE)]}"/></a:solidFill>'
        f"</c:spPr></c:dPt>"
        for i in range(n)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:pieChart>
        <c:varyColors val="1"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          {color_pts}
          <c:dLbls>
            <c:numFmt formatCode="0%" sourceLinked="0"/>
            <c:showLegendKey val="0"/><c:showVal val="0"/>
            <c:showCatName val="1"/><c:showSerName val="0"/>
            <c:showPercent val="1"/><c:showBubbleSize val="0"/>
            <c:separator>: </c:separator>
          </c:dLbls>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
        </c:ser>
        <c:firstSliceAng val="0"/>
      </c:pieChart>
    </c:plotArea>
    <c:legend><c:legendPos val="r"/><c:overlay val="0"/></c:legend>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _build_donut_chart_xml(
    labels: list[str],
    values: list[int],
    title: str,
    hole_size: int = 50,
) -> str:
    """Generate DrawingML XML for a doughnut chart (modern variant of pie)."""
    n = len(labels)
    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(l))}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    pts_val = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
        for i, v in enumerate(values)
    )
    color_pts = "".join(
        f'<c:dPt><c:idx val="{i}"/><c:spPr>'
        f'<a:solidFill><a:srgbClr val="{_CHART_PALETTE[i % len(_CHART_PALETTE)]}"/></a:solidFill>'
        f"</c:spPr></c:dPt>"
        for i in range(n)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:doughnutChart>
        <c:varyColors val="1"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          {color_pts}
          <c:dLbls>
            <c:numFmt formatCode="0%" sourceLinked="0"/>
            <c:showLegendKey val="0"/><c:showVal val="0"/>
            <c:showCatName val="1"/><c:showSerName val="0"/>
            <c:showPercent val="1"/><c:showBubbleSize val="0"/>
            <c:separator>: </c:separator>
          </c:dLbls>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
        </c:ser>
        <c:firstSliceAng val="0"/>
        <c:holeSize val="{hole_size}"/>
      </c:doughnutChart>
    </c:plotArea>
    <c:legend><c:legendPos val="r"/><c:overlay val="0"/></c:legend>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _build_line_chart_xml(
    labels: list[str],
    values: list[int],
    title: str,
) -> str:
    """Generate DrawingML XML for a line chart with markers — ideal for timelines."""
    n = len(labels)
    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(l))}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    pts_val = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
        for i, v in enumerate(values)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:lineChart>
        <c:barDir val="col"/>
        <c:grouping val="standard"/>
        <c:varyColors val="0"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          <c:spPr>
            <a:ln w="25400">
              <a:solidFill><a:srgbClr val="0D7F7A"/></a:solidFill>
            </a:ln>
          </c:spPr>
          <c:marker>
            <c:symbol val="circle"/>
            <c:size val="5"/>
            <c:spPr>
              <a:solidFill><a:srgbClr val="0D7F7A"/></a:solidFill>
              <a:ln><a:solidFill><a:srgbClr val="0A5F5B"/></a:solidFill></a:ln>
            </c:spPr>
          </c:marker>
          <c:dLbls>
            <c:numFmt formatCode="General" sourceLinked="0"/>
            <c:showLegendKey val="0"/><c:showVal val="1"/>
            <c:showCatName val="0"/><c:showSerName val="0"/>
            <c:showPercent val="0"/><c:showBubbleSize val="0"/>
          </c:dLbls>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
          <c:smooth val="0"/>
        </c:ser>
        <c:axId val="200"/><c:axId val="201"/>
      </c:lineChart>
      <c:catAx>
        <c:axId val="200"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="b"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="201"/>
      </c:catAx>
      <c:valAx>
        <c:axId val="201"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="l"/>
        <c:numFmt formatCode="General" sourceLinked="0"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="200"/>
        <c:crossBetween val="between"/>
      </c:valAx>
    </c:plotArea>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _build_area_chart_xml(
    labels: list[str],
    values: list[int],
    title: str,
) -> str:
    """Generate DrawingML XML for an area chart — ideal for cumulative volume over time."""
    n = len(labels)
    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(l))}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    pts_val = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
        for i, v in enumerate(values)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:areaChart>
        <c:grouping val="standard"/>
        <c:varyColors val="0"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          <c:spPr>
            <a:solidFill>
              <a:srgbClr val="0D7F7A">
                <a:alpha val="60000"/>
              </a:srgbClr>
            </a:solidFill>
            <a:ln w="19050">
              <a:solidFill><a:srgbClr val="0A5F5B"/></a:solidFill>
            </a:ln>
          </c:spPr>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
        </c:ser>
        <c:axId val="300"/><c:axId val="301"/>
      </c:areaChart>
      <c:catAx>
        <c:axId val="300"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="b"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="301"/>
      </c:catAx>
      <c:valAx>
        <c:axId val="301"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="l"/>
        <c:numFmt formatCode="General" sourceLinked="0"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="300"/>
        <c:crossBetween val="between"/>
      </c:valAx>
    </c:plotArea>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _build_stacked_bar_chart_xml(
    categories: list[str],
    series_names: list[str],
    series_values_list: list[list[int]],
    title: str,
    percent_stacked: bool = False,
    horizontal: bool = True,
    series_colors: list[str] | None = None,
) -> str:
    """Generate DrawingML XML for a stacked (or 100% stacked) multi-series bar chart.

    Args:
        categories: Category labels (X axis for vertical bars, Y axis for horizontal).
        series_names: One name per data series.
        series_values_list: Parallel list of value arrays — one list per series.
        title: Chart title.
        percent_stacked: Use 100% stacked grouping instead of absolute stacked.
        horizontal: True = horizontal bars; False = vertical (column) bars.
        series_colors: Optional explicit hex colors per series; defaults to _CHART_PALETTE.
    """
    n_cats = len(categories)
    grouping = "percentStacked" if percent_stacked else "stacked"
    bar_dir = "bar" if horizontal else "col"
    cat_pos = "l" if horizontal else "b"
    val_pos = "b" if horizontal else "l"
    orientation = "maxMin" if horizontal else "minMax"

    pts_cat = "".join(
        f'<c:pt idx="{i}"><c:v>{_xe(str(c))}</c:v></c:pt>'
        for i, c in enumerate(categories)
    )

    series_xml_parts = []
    for s_idx, (s_name, s_values) in enumerate(zip(series_names, series_values_list)):
        color = (series_colors[s_idx] if series_colors and s_idx < len(series_colors)
                 else _CHART_PALETTE[s_idx % len(_CHART_PALETTE)])
        pts_val = "".join(
            f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>'
            for i, v in enumerate(s_values)
        )
        fmt_code = "0%" if percent_stacked else "General"
        series_xml_parts.append(f"""
        <c:ser>
          <c:idx val="{s_idx}"/><c:order val="{s_idx}"/>
          <c:tx><c:strRef><c:strCache>
            <c:ptCount val="1"/>
            <c:pt idx="0"><c:v>{_xe(s_name)}</c:v></c:pt>
          </c:strCache></c:strRef></c:tx>
          <c:spPr>
            <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
            <a:ln><a:solidFill><a:srgbClr val="{color}"/></a:solidFill></a:ln>
          </c:spPr>
          <c:dLbls>
            <c:numFmt formatCode="{fmt_code}" sourceLinked="0"/>
            <c:showLegendKey val="0"/><c:showVal val="1"/>
            <c:showCatName val="0"/><c:showSerName val="0"/>
            <c:showPercent val="0"/><c:showBubbleSize val="0"/>
          </c:dLbls>
          <c:cat><c:strRef><c:strCache>
            <c:ptCount val="{n_cats}"/>{pts_cat}
          </c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:numCache>
            <c:formatCode>General</c:formatCode>
            <c:ptCount val="{n_cats}"/>{pts_val}
          </c:numCache></c:numRef></c:val>
        </c:ser>""")

    all_series = "".join(series_xml_parts)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{_NS_C}" xmlns:a="{_NS_A}" xmlns:r="{_NS_R}">
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/>
        <a:p><a:pPr><a:defRPr b="1" sz="1100"/></a:pPr>
          <a:r><a:rPr lang="en-US" b="1"/><a:t>{_xe(title)}</a:t></a:r>
        </a:p>
      </c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:autoTitleDeleted val="0"/>
    <c:plotArea>
      <c:barChart>
        <c:barDir val="{bar_dir}"/>
        <c:grouping val="{grouping}"/>
        <c:varyColors val="0"/>
        {all_series}
        <c:axId val="400"/><c:axId val="401"/>
      </c:barChart>
      <c:catAx>
        <c:axId val="400"/>
        <c:scaling><c:orientation val="{orientation}"/></c:scaling>
        <c:delete val="0"/><c:axPos val="{cat_pos}"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="401"/>
      </c:catAx>
      <c:valAx>
        <c:axId val="401"/>
        <c:scaling><c:orientation val="minMax"/></c:scaling>
        <c:delete val="0"/><c:axPos val="{val_pos}"/>
        <c:numFmt formatCode="General" sourceLinked="0"/>
        <c:tickLblPos val="nextTo"/><c:crossAx val="400"/>
        <c:crossBetween val="between"/>
      </c:valAx>
    </c:plotArea>
    <c:legend><c:legendPos val="b"/><c:overlay val="0"/></c:legend>
    <c:plotVisOnly val="1"/>
  </c:chart>
</c:chartSpace>"""


def _insert_native_chart(
    document: Document,
    chart_xml: str,
    width_inches: float = 6.0,
    height_inches: float = 3.2,
) -> None:
    """Embed a DrawingML native Office chart into a python-docx Document."""
    chart_count = sum(1 for rel in document.part.rels.values() if _CHART_REL in rel.reltype)
    chart_idx = chart_count + 1

    chart_part = _Part(
        partname=_PackURI(f"/word/charts/chart{chart_idx}.xml"),
        content_type=_CHART_CT,
        blob=chart_xml.encode("utf-8"),
        package=document.part.package,
    )
    rel_id = document.part.relate_to(chart_part, _CHART_REL)

    cx, cy = _emu(width_inches), _emu(height_inches)
    drawing_xml = (
        f'<w:drawing xmlns:w="{_NS_W}">'
        f'<wp:inline xmlns:wp="{_NS_WP}" distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{1000 + chart_idx}" name="Chart {chart_idx}"/>'
        f"<wp:cNvGraphicFramePr/>"
        f'<a:graphic xmlns:a="{_NS_A}">'
        f'<a:graphicData uri="{_NS_C}">'
        f'<c:chart xmlns:c="{_NS_C}" xmlns:r="{_NS_R}" r:id="{rel_id}"/>'
        f"</a:graphicData>"
        f"</a:graphic>"
        f"</wp:inline>"
        f"</w:drawing>"
    )
    para = document.add_paragraph()
    run = para.add_run()
    run._r.append(_etree.fromstring(drawing_xml))


def ensure_default_top_ranking_template(template_path: str | Path) -> Path:
    """Create the default importable DOCX template when it is missing."""

    resolved_template_path = Path(template_path)
    if resolved_template_path.exists() and not _template_contains_legacy_text(
        resolved_template_path
    ):
        return resolved_template_path

    _create_default_top_ranking_template(resolved_template_path)
    return resolved_template_path


def _template_contains_legacy_text(template_path: Path) -> bool:
    """Return True when the bundled default template still has old dev copy."""

    try:
        document = Document(template_path)
    except Exception:
        return True

    legacy_markers = (
        "Ranking tables are inserted after template rendering",
        "Chart image supplied.",
        "Field coverage is inserted after template rendering.",
    )
    template_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return any(marker in template_text for marker in legacy_markers)


def _create_default_top_ranking_template(template_path: Path) -> None:
    """Write the default clean DOCX template."""

    resolved_template_path = template_path
    resolved_template_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    core = document.core_properties
    core.title = "GTI Top Targets Ranking Template"
    core.subject = "docxtpl template for GTI Top Rankings exports"

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)

    document.add_heading("{{ report_title }}", level=0)
    document.add_paragraph("Period: {{ period }}")
    document.add_paragraph("Generated at: {{ generated_at }}")
    document.add_paragraph("Scope / query: {{ query_used }}")
    document.add_paragraph("{{ preview_mode_note }}")

    document.add_page_break()
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph("{{ executive_summary }}")
    document.add_paragraph("Selected rankings: {{ selected_rankings_text }}")
    document.add_paragraph("Main top results: {{ main_top_results }}")
    document.add_paragraph("Field coverage: {{ field_coverage_summary }}")

    document.add_heading("Methodology", level=1)
    document.add_paragraph("GTI query used: {{ query_used }}")
    document.add_paragraph("{{ preview_only_explanation }}")
    document.add_paragraph("Request estimate: {{ estimated_api_requests }} estimated API request(s).")
    document.add_paragraph("Actual Intelligence Search requests: {{ actual_search_requests }}.")
    document.add_paragraph("{{ methodology }}")

    document.add_heading("Rankings", level=1)
    document.add_paragraph("The ranking tables and charts below use already computed preview fields.")

    document.add_heading("Limitations", level=1)
    document.add_paragraph("Counts represent GTI collections, not confirmed incident counts.", style=None)
    document.add_paragraph("Preview-only fields may be incomplete.", style=None)
    document.add_paragraph("Crowdsourced collections may introduce noise.", style=None)

    document.add_heading("Appendix", level=1)
    document.add_paragraph("Detailed field coverage is included below.")
    document.add_paragraph("{{ technical_debug_note }}")

    document.save(resolved_template_path)


def generate_top_ranking_docx(
    ranking_result: dict,
    template_path: str,
    output_path: str,
) -> str:
    """Render a GTI Top Rankings DOCX report from an existing result object."""

    resolved_template = ensure_default_top_ranking_template(template_path)
    resolved_output = Path(output_path)
    resolved_output.parent.mkdir(parents=True, exist_ok=True)

    sanitized_result = _sanitize_ranking_result(ranking_result)
    include_debug = bool(sanitized_result.get("include_technical_debug"))
    chart_temp_paths = _write_chart_images(sanitized_result.get("charts", {}))

    template = DocxTemplate(str(resolved_template))
    context = _build_docx_context(
        sanitized_result,
        include_debug=include_debug,
        chart_temp_paths=chart_temp_paths,
        template=template,
    )
    template.render(context)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        rendered_path = Path(temp_file.name)
    template.save(rendered_path)

    document = Document(rendered_path)
    _append_ranking_tables(document, context)
    _append_cross_analysis(document, context)
    _append_field_coverage(document, context)
    _append_ttp_diagnostics(document, context)
    _append_optional_debug(document, context, include_debug)

    document.save(resolved_output)
    rendered_path.unlink(missing_ok=True)
    for chart_path in chart_temp_paths.values():
        chart_path.unlink(missing_ok=True)

    return str(resolved_output)


def _sanitize_ranking_result(ranking_result: dict[str, Any]) -> dict[str, Any]:
    """Remove sensitive or overly raw fields from the export input."""

    sanitized = dict(ranking_result or {})
    for key in (
        "api_key",
        "x_api_key",
        "raw_data",
        "raw_json",
    ):
        sanitized.pop(key, None)
    return sanitized


def _build_docx_context(
    ranking_result: dict[str, Any],
    include_debug: bool,
    chart_temp_paths: dict[str, Path],
    template: DocxTemplate,
) -> dict[str, Any]:
    """Build a docxtpl-safe context from the Top Rankings response."""

    rankings = ranking_result.get("rankings") if isinstance(ranking_result.get("rankings"), dict) else {}
    fields_coverage = ranking_result.get("fields_coverage") if isinstance(ranking_result.get("fields_coverage"), dict) else {}
    collections_analyzed = _safe_int(ranking_result.get("collections_analyzed"))
    selected_rankings = [
        str(item)
        for item in ranking_result.get("selected_rankings", [])
        if item is not None
    ]

    ranking_tables = {
        template_key: _normalize_ranking_rows(rankings.get(result_key, []))
        for template_key, result_key in RANKING_RESULT_KEYS.items()
    }
    ranking_tables["top_tactics"] = _normalize_ranking_rows(
        ranking_result.get("top_tactics", [])
    )
    ranking_tables["top_techniques"] = _normalize_ranking_rows(
        ranking_result.get("top_techniques", [])
    )
    ranking_tables["top_subtechniques"] = _normalize_ranking_rows(
        ranking_result.get("top_subtechniques", [])
    )
    preview_collections = ranking_result.get("collection_preview_fields", [])
    technical_debug = (
        ranking_result.get("technical_debug", {})
        if isinstance(ranking_result.get("technical_debug"), dict)
        else {}
    )
    ranking_debug = (
        technical_debug.get("ranking_debug", {})
        if isinstance(technical_debug.get("ranking_debug"), dict)
        else {}
    )
    threat_categories = _build_single_field_ranking(
        preview_collections,
        "threat_categories",
    )
    if threat_categories:
        ranking_tables["threat_categories"] = threat_categories
    cross_analysis = build_cross_analysis_matrices(preview_collections)
    ranking_notes = {
        f"{key}_note": "" if rows else EMPTY_RANKING_NOTE
        for key, rows in ranking_tables.items()
    }

    context: dict[str, Any] = {
        "report_title": "GTI Top Targets Ranking",
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "period": str(ranking_result.get("period") or ""),
        "start_year": ranking_result.get("start_year"),
        "month": ranking_result.get("month"),
        "query_used": str(ranking_result.get("query_used") or ""),
        "collections_analyzed": collections_analyzed,
        "collections_seen": _safe_int(ranking_result.get("collections_seen")),
        "max_collections": _safe_int(ranking_result.get("max_collections")),
        "pages_fetched": _safe_int(ranking_result.get("pages_fetched")),
        "actual_search_requests": _safe_int(ranking_result.get("actual_search_requests")),
        "estimated_api_requests": _safe_int(ranking_result.get("estimated_api_requests")),
        "selected_rankings": selected_rankings,
        "selected_rankings_text": ", ".join(selected_rankings) if selected_rankings else "None",
        "fields_coverage": fields_coverage,
        "field_coverage_summary": _build_field_coverage_summary(
            fields_coverage,
            collections_analyzed,
        ),
        "methodology": str(ranking_result.get("methodology") or ""),
        "preview_mode_note": "Preview-only mode: uses only fields returned by GTI Intelligence Search.",
        "preview_only_explanation": (
            "This report uses Intelligence Search preview fields and avoids expensive "
            "per-collection detail calls."
        ),
        "executive_summary": (
            f"Analyzed {collections_analyzed} GTI collections for "
            f"{ranking_result.get('period') or 'the selected period'}."
        ),
        "main_top_results": _build_main_top_results(ranking_tables),
        "technical_debug_note": (
            "Technical debug appendix included."
            if include_debug
            else "Technical debug appendix was not included."
        ),
        "ttp_analysis": (
            ranking_result.get("ttp_analysis", {})
            if isinstance(ranking_result.get("ttp_analysis"), dict)
            else {}
        ),
        "debug_attribute_keys_frequency": (
            ranking_result.get(
                "debug_attribute_keys_frequency",
                ranking_debug.get("debug_attribute_keys_frequency", {}),
            )
            if include_debug
            else {}
        ),
        "debug_sample_collection_fields": (
            ranking_result.get(
                "debug_sample_collection_fields",
                ranking_debug.get("debug_sample_collection_fields", []),
            )
            if include_debug
            else []
        ),
        "_chart_temp_paths": chart_temp_paths,
        "cross_analysis": cross_analysis,
    }
    context.update(ranking_tables)
    context.update(ranking_notes)
    context.update(_build_chart_context(template, chart_temp_paths))
    return context


def _normalize_ranking_rows(rows: Any) -> list[dict[str, Any]]:
    """Normalize ranking rows for tables."""

    if not isinstance(rows, list):
        return []

    normalized_rows: list[dict[str, Any]] = []
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            continue
        normalized_rows.append(
            {
                "rank": _safe_int(row.get("rank"), index),
                "name": str(row.get("name") or "Unknown"),
                "collection_count": _safe_int(
                    row.get("collection_count", row.get("report_count"))
                ),
            }
        )
    return normalized_rows


def _build_single_field_ranking(
    collections: Any,
    field_name: str,
    top_n: int = 25,
) -> list[dict[str, Any]]:
    """Build a simple distinct-per-collection ranking from preview fields."""

    if not isinstance(collections, list):
        return []

    counter: dict[str, int] = {}
    display: dict[str, str] = {}
    for collection in collections:
        if not isinstance(collection, dict):
            continue
        values = _extract_docx_names(collection.get(field_name))
        for value in set(values):
            normalized = value.casefold()
            display.setdefault(normalized, value)
            counter[normalized] = counter.get(normalized, 0) + 1

    return [
        {
            "rank": index + 1,
            "name": display[key],
            "collection_count": count,
        }
        for index, (key, count) in enumerate(
            sorted(counter.items(), key=lambda item: (-item[1], display[item[0]].casefold()))[:top_n]
        )
    ]


def build_cross_analysis_matrices(
    collections: Any,
    top_rows: int = 8,
    top_columns: int = 8,
) -> dict[str, Any]:
    """Build co-occurrence matrices from collection preview fields only."""

    if not isinstance(collections, list):
        collections = []

    matrix_specs = {
        "industries_by_tags": ("targeted_industries", "tags"),
        "industries_by_collection_type": ("targeted_industries", "collection_type"),
        "industries_by_targeted_region": ("targeted_industries", "targeted_regions"),
        "timeline_by_collection_type": ("timeline", "collection_type"),
        "source_region_by_targeted_region": ("source_regions", "targeted_regions"),
    }

    return {
        matrix_key: _build_cooccurrence_matrix(
            collections=collections,
            row_field=row_field,
            column_field=column_field,
            top_rows=top_rows,
            top_columns=top_columns,
        )
        for matrix_key, (row_field, column_field) in matrix_specs.items()
    }


def _build_cooccurrence_matrix(
    collections: list[Any],
    row_field: str,
    column_field: str,
    top_rows: int,
    top_columns: int,
) -> dict[str, Any]:
    """Build one distinct-per-collection co-occurrence matrix."""

    pair_counter: dict[tuple[str, str], int] = {}
    row_counter: dict[str, int] = {}
    column_counter: dict[str, int] = {}
    row_display: dict[str, str] = {}
    column_display: dict[str, str] = {}
    eligible_collections = 0

    for collection in collections:
        if not isinstance(collection, dict):
            continue
        row_values = _extract_matrix_values(collection, row_field)
        column_values = _extract_matrix_values(collection, column_field)
        if not row_values or not column_values:
            continue

        eligible_collections += 1
        normalized_rows = {_normalize_matrix_value(value): value for value in row_values}
        normalized_columns = {_normalize_matrix_value(value): value for value in column_values}
        for row_key, row_label in normalized_rows.items():
            if not row_key:
                continue
            row_display.setdefault(row_key, row_label)
            row_counter[row_key] = row_counter.get(row_key, 0) + 1
        for column_key, column_label in normalized_columns.items():
            if not column_key:
                continue
            column_display.setdefault(column_key, column_label)
            column_counter[column_key] = column_counter.get(column_key, 0) + 1
        for row_key in normalized_rows:
            for column_key in normalized_columns:
                if not row_key or not column_key:
                    continue
                pair_counter[(row_key, column_key)] = pair_counter.get((row_key, column_key), 0) + 1

    selected_rows = [
        key
        for key, _ in sorted(
            row_counter.items(),
            key=lambda item: (-item[1], row_display[item[0]].casefold()),
        )[:top_rows]
    ]
    selected_columns = [
        key
        for key, _ in sorted(
            column_counter.items(),
            key=lambda item: (-item[1], column_display[item[0]].casefold()),
        )[:top_columns]
    ]
    table_rows = []
    for row_key in selected_rows:
        cells = [
            pair_counter.get((row_key, column_key), 0)
            for column_key in selected_columns
        ]
        table_rows.append(
            {
                "label": row_display[row_key],
                "cells": cells,
            }
        )

    top_cells = [
        {
            "row": row_display[row_key],
            "column": column_display[column_key],
            "count": count,
        }
        for (row_key, column_key), count in pair_counter.items()
    ]
    top_cells = sorted(
        top_cells,
        key=lambda item: (-item["count"], item["row"].casefold(), item["column"].casefold()),
    )[:5]

    return {
        "eligible_collections": eligible_collections,
        "columns": [column_display[key] for key in selected_columns],
        "rows": table_rows,
        "top_cells": top_cells,
        "interpretation": _build_cross_analysis_interpretation(top_cells, eligible_collections),
    }


def _extract_matrix_values(collection: dict[str, Any], field_name: str) -> list[str]:
    """Extract field values for cross-analysis matrices."""

    if field_name == "timeline":
        bucket = _build_docx_timeline_bucket(collection.get("creation_date"))
        return [bucket] if bucket else []

    return _extract_docx_names(collection.get(field_name))


def _extract_docx_names(value: Any) -> list[str]:
    """Extract readable values from preview fields for reporting."""

    if value is None:
        return []
    if isinstance(value, str):
        stripped = value.strip()
        return [stripped] if stripped else []
    if isinstance(value, bool):
        return []
    if isinstance(value, (int, float)):
        return [str(value)]
    if isinstance(value, list):
        names: list[str] = []
        for item in value:
            names.extend(_extract_docx_names(item))
        return _dedupe_names(names)
    if isinstance(value, dict):
        names: list[str] = []
        for key in ("name", "label", "title", "value", "id"):
            if key in value:
                names.extend(_extract_docx_names(value.get(key)))
                break
        for key, nested_value in value.items():
            if key in ("name", "label", "title", "value", "id"):
                continue
            names.extend(_extract_docx_names(nested_value))
        return _dedupe_names(names)

    return []


def _dedupe_names(values: list[str]) -> list[str]:
    """Dedupe extracted names while preserving order."""

    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        normalized = value.casefold()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(value)
    return deduped


def _normalize_matrix_value(value: str) -> str:
    """Normalize a matrix label for distinct-per-collection counting."""

    return " ".join(str(value).split()).casefold()


def _build_docx_timeline_bucket(value: Any) -> str | None:
    """Build a YYYY-MM bucket for cross-analysis timeline rows."""

    if value is None:
        return None
    text = str(value).strip()
    match = re.match(r"^(\d{4})-(\d{2})", text)
    if match:
        return f"{match.group(1)}-{match.group(2)}"
    if re.fullmatch(r"\d+(?:\.\d+)?", text):
        try:
            timestamp = float(text)
            if timestamp > 1_000_000_000_000:
                timestamp = timestamp / 1000
            return datetime.fromtimestamp(timestamp, tz=timezone.utc).strftime("%Y-%m")
        except (OSError, OverflowError, ValueError):
            return None
    return None


def _build_cross_analysis_interpretation(
    top_cells: list[dict[str, Any]],
    eligible_collections: int,
) -> str:
    """Generate a concise interpretation from top co-occurrence cells."""

    if not top_cells:
        return (
            "Not enough overlapping preview metadata was present to build this matrix."
        )

    strongest = top_cells[0]
    return (
        f"The strongest metadata co-occurrence is {strongest['row']} x "
        f"{strongest['column']} with {strongest['count']} GTI collection(s) "
        f"among {eligible_collections} eligible collection(s). These counts reflect "
        "GTI collection metadata, not confirmed incident counts."
    )


def _build_field_coverage_summary(
    fields_coverage: dict[str, Any],
    collections_analyzed: int,
) -> str:
    """Create a compact field coverage sentence for the executive summary."""

    if not fields_coverage:
        return "Field coverage was not available."

    labels = {
        "targeted_industries": "targeted industries",
        "targeted_regions": "targeted regions",
        "source_regions": "source regions",
        "tags": "tags / themes",
        "collection_type": "collection type",
        "timeline": "timeline",
        "targeted_organizations": "targeted organizations",
    }
    parts = [
        f"{label}: {_safe_int(fields_coverage.get(key))}/{collections_analyzed}"
        for key, label in labels.items()
    ]
    return "; ".join(parts)


def _build_main_top_results(ranking_tables: dict[str, list[dict[str, Any]]]) -> str:
    """Summarize the first row of each available ranking."""

    summary_items = []
    for ranking_key, ranking_label in RANKING_LABELS.items():
        rows = ranking_tables.get(ranking_key, [])
        if not rows:
            continue
        first_row = rows[0]
        summary_items.append(
            f"{ranking_label}: {first_row['name']} ({first_row['collection_count']} collections)"
        )

    return "; ".join(summary_items) if summary_items else EMPTY_RANKING_NOTE


def _build_chart_context(
    template: DocxTemplate,
    chart_temp_paths: dict[str, Path],
) -> dict[str, Any]:
    """Prepare chart notes/placeholders for docxtpl templates."""

    chart_keys = (
        "industry_chart",
        "targeted_regions_chart",
        "source_regions_chart",
        "tags_chart",
        "collection_type_chart",
        "timeline_chart",
    )
    context: dict[str, Any] = {}
    for chart_key in chart_keys:
        context[f"{chart_key}_note"] = (
            "Chart image supplied."
            if chart_key in chart_temp_paths
            else "Chart image was not supplied by the app; table is included instead."
        )
        context[chart_key] = (
            InlineImage(template, str(chart_temp_paths[chart_key]), width=Inches(6.2))
            if chart_key in chart_temp_paths
            else ""
        )
    return context


def _write_chart_images(charts: Any) -> dict[str, Path]:
    """Decode optional client-supplied PNG data URLs into temporary image files."""

    if not isinstance(charts, dict):
        return {}

    chart_paths: dict[str, Path] = {}
    for chart_key, chart_data in charts.items():
        if not isinstance(chart_data, str) or not chart_data.startswith("data:image/png;base64,"):
            continue
        # Chart support is intentionally tolerant. The default app currently does
        # not generate PNG charts, so a malformed or missing chart must not fail export.
        try:
            import base64

            png_bytes = base64.b64decode(chart_data.split(",", 1)[1], validate=True)
        except Exception:
            continue
        temp_path = Path(tempfile.gettempdir()) / f"gti_{_slugify(str(chart_key))}.png"
        temp_path.write_bytes(png_bytes)
        chart_paths[str(chart_key)] = temp_path

    return chart_paths


def _append_ranking_tables(document: Document, context: dict[str, Any]) -> None:
    """Append ranking tables and native charts to the rendered report document."""

    document.add_page_break()
    document.add_heading("Rankings", level=1)
    for ranking_key, ranking_label in RANKING_LABELS.items():
        rows = context.get(ranking_key, [])
        document.add_heading(ranking_label, level=2)
        if not rows:
            document.add_paragraph(context.get(f"{ranking_key}_note") or EMPTY_RANKING_NOTE)
            continue

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "#"
        header_cells[1].text = "Name"
        header_cells[2].text = "Collections"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row["rank"])
            cells[1].text = str(row["name"])
            cells[2].text = str(row["collection_count"])

        chart_rows = rows[:10]
        labels = [str(r["name"]) for r in chart_rows]
        values = [_safe_int(r["collection_count"]) for r in chart_rows]
        try:
            if ranking_key == "collection_type_distribution":
                # Pie chart
                xml = _build_pie_chart_xml(labels, values, ranking_label)
                _insert_native_chart(document, xml, width_inches=5.0, height_inches=3.2)
                # Donut chart — modern alternative
                xml = _build_donut_chart_xml(labels, values, ranking_label + " — Donut")
                _insert_native_chart(document, xml, width_inches=5.0, height_inches=3.2)
                # Horizontal bar for comparison
                xml = _build_bar_chart_xml(labels, values, ranking_label + " — Bars", horizontal=True)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=max(2.0, len(labels) * 0.32 + 0.8))
            elif ranking_key == "timeline":
                # Column (vertical bar) chart
                xml = _build_bar_chart_xml(labels, values, ranking_label + " — Columns", horizontal=False)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=3.2)
                # Line chart — trend view
                xml = _build_line_chart_xml(labels, values, ranking_label + " — Line")
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=3.2)
                # Area chart — cumulative volume view
                xml = _build_area_chart_xml(labels, values, ranking_label + " — Area")
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=3.2)
            elif ranking_key in ("top_tactics", "top_techniques", "top_subtechniques"):
                height = max(2.0, min(4.0, len(labels) * 0.32 + 0.8))
                # Horizontal bar
                xml = _build_bar_chart_xml(labels, values, ranking_label, horizontal=True)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=height)
                # Vertical column chart
                xml = _build_bar_chart_xml(labels, values, ranking_label + " — Columns", horizontal=False)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=3.2)
                # Pie chart — relative share
                xml = _build_pie_chart_xml(labels, values, ranking_label + " — Pie")
                _insert_native_chart(document, xml, width_inches=5.0, height_inches=3.2)
                # Donut chart — modern alternative
                xml = _build_donut_chart_xml(labels, values, ranking_label + " — Donut")
                _insert_native_chart(document, xml, width_inches=5.0, height_inches=3.2)
            else:
                height = max(2.0, min(4.0, len(labels) * 0.32 + 0.8))
                # Horizontal bar (primary)
                xml = _build_bar_chart_xml(labels, values, ranking_label, horizontal=True)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=height)
                # Vertical column chart
                xml = _build_bar_chart_xml(labels, values, ranking_label + " — Columns", horizontal=False)
                _insert_native_chart(document, xml, width_inches=6.0, height_inches=3.2)
        except Exception:
            pass



def _append_cross_analysis(document: Document, context: dict[str, Any]) -> None:
    """Append cross-analysis matrices after basic rankings."""

    document.add_heading("Cross-analysis", level=1)
    document.add_paragraph(
        "These matrices count co-occurring GTI collection metadata values. Counts "
        "represent GTI collections, not confirmed incident counts."
    )
    matrices = context.get("cross_analysis", {})
    if not isinstance(matrices, dict) or not matrices:
        document.add_paragraph("Cross-analysis was not available for this export.")
        return

    for matrix_key, matrix in matrices.items():
        label = CROSS_ANALYSIS_LABELS.get(matrix_key, matrix_key)
        document.add_heading(label, level=2)
        eligible = _safe_int(matrix.get("eligible_collections") if isinstance(matrix, dict) else 0)
        document.add_paragraph(f"Eligible collections: {eligible}")
        document.add_paragraph(
            str(matrix.get("interpretation") or "No interpretation available.")
            if isinstance(matrix, dict)
            else "No interpretation available."
        )

        columns = matrix.get("columns", []) if isinstance(matrix, dict) else []
        rows = matrix.get("rows", []) if isinstance(matrix, dict) else []
        if not columns or not rows:
            document.add_paragraph("Not enough overlapping preview fields to build this matrix.")
            continue

        max_value = max(
            [max(row.get("cells", [0]) or [0]) for row in rows if isinstance(row, dict)]
            or [0]
        )
        table = document.add_table(rows=1, cols=len(columns) + 1)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = ""
        for index, column in enumerate(columns, start=1):
            header_cells[index].text = str(column)

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("label", ""))
            for index, value in enumerate(row.get("cells", []), start=1):
                cells[index].text = str(value)
                _shade_cell(cells[index], _heatmap_shade(_safe_int(value), max_value))


def _heatmap_shade(value: int, max_value: int) -> str:
    """Return a light teal heatmap color for a matrix cell."""

    if value <= 0 or max_value <= 0:
        return "FFFFFF"
    intensity = min(1.0, value / max_value)
    # Blend white toward the product teal color.
    base = (13, 127, 122)
    blended = tuple(round(255 - (255 - channel) * intensity * 0.55) for channel in base)
    return "".join(f"{channel:02X}" for channel in blended)


def _shade_cell(cell: Any, fill: str) -> None:
    """Apply a background fill color to a Word table cell."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)
    except Exception:
        return


def _append_field_coverage(document: Document, context: dict[str, Any]) -> None:
    """Append field coverage diagnostics with a stacked bar chart."""

    document.add_heading("Appendix: Field Coverage", level=1)
    coverage = context.get("fields_coverage", {})
    total = _safe_int(context.get("collections_analyzed"))
    coverage_fields = (
        "targeted_industries",
        "targeted_regions",
        "source_regions",
        "tags",
        "collection_type",
        "timeline",
        "targeted_organizations",
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Collections with data"
    header_cells[2].text = "Collections analyzed"
    for field_name in coverage_fields:
        cells = table.add_row().cells
        cells[0].text = field_name
        cells[1].text = str(_safe_int(coverage.get(field_name)))
        cells[2].text = str(total)

    # 100% stacked bar chart — visual coverage overview
    if total > 0:
        try:
            cats = [f.replace("_", " ").title() for f in coverage_fields]
            with_data = [_safe_int(coverage.get(f)) for f in coverage_fields]
            without_data = [max(0, total - v) for v in with_data]
            xml = _build_stacked_bar_chart_xml(
                categories=cats,
                series_names=["With data", "Without data"],
                series_values_list=[with_data, without_data],
                title="Field Coverage — 100% Stacked",
                percent_stacked=True,
                horizontal=True,
                series_colors=["0D7F7A", "D9D9D9"],
            )
            height = max(2.5, min(4.5, len(cats) * 0.38 + 0.8))
            _insert_native_chart(document, xml, width_inches=6.0, height_inches=height)
            # Absolute stacked bar chart — absolute counts
            xml_abs = _build_stacked_bar_chart_xml(
                categories=cats,
                series_names=["With data", "Without data"],
                series_values_list=[with_data, without_data],
                title="Field Coverage — Absolute Counts",
                percent_stacked=False,
                horizontal=True,
                series_colors=["0D7F7A", "D9D9D9"],
            )
            _insert_native_chart(document, xml_abs, width_inches=6.0, height_inches=height)
        except Exception:
            pass


def _append_ttp_diagnostics(document: Document, context: dict[str, Any]) -> None:
    """Append hard TTP diagnostic fields used by the app."""

    ttp = context.get("ttp_analysis", {})
    if not isinstance(ttp, dict) or not ttp:
        return

    first_debug = ttp.get("ttp_first_successful_debug", {})
    if not isinstance(first_debug, dict):
        first_debug = {}

    document.add_heading("Appendix: TTP Diagnostics", level=1)
    if ttp.get("warning_message"):
        document.add_paragraph(str(ttp.get("warning_message")))

    rows = (
        ("ttp_lookups_attempted", ttp.get("ttp_lookups_attempted", 0)),
        ("ttp_lookups_succeeded", ttp.get("ttp_lookups_succeeded", 0)),
        ("ttp_eligible_collections", ttp.get("ttp_eligible_collections", 0)),
        (
            "ttp_first_successful_collection_id",
            ttp.get("ttp_first_successful_collection_id", ""),
        ),
        (
            "ttp_first_successful_debug.tactics_count",
            first_debug.get("tactics_count", 0),
        ),
        ("top_tactics count", len(context.get("top_tactics", []))),
        ("top_techniques count", len(context.get("top_techniques", []))),
        ("top_subtechniques count", len(context.get("top_subtechniques", []))),
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Diagnostic"
    header_cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)

    document.add_heading("ttp_lookup_attempt_samples", level=2)
    samples = ttp.get("ttp_lookup_attempt_samples", [])
    if isinstance(samples, list) and samples:
        document.add_paragraph(json.dumps(samples, indent=2, default=str))
    else:
        document.add_paragraph("No TTP lookup samples were recorded.")


def _append_optional_debug(
    document: Document,
    context: dict[str, Any],
    include_debug: bool,
) -> None:
    """Append technical debug data only when explicitly requested."""

    if not include_debug:
        return

    document.add_heading("Technical Debug Appendix", level=1)
    document.add_heading("Attribute Key Frequency", level=2)
    frequency = context.get("debug_attribute_keys_frequency", {})
    if isinstance(frequency, dict) and frequency:
        for key, count in sorted(frequency.items(), key=lambda item: (-_safe_int(item[1]), str(item[0]))):
            document.add_paragraph(f"{key}: {count}")
    else:
        document.add_paragraph("No attribute key diagnostics were available.")

    document.add_heading("Sample Collection Fields", level=2)
    samples = context.get("debug_sample_collection_fields", [])
    if isinstance(samples, list) and samples:
        for sample in samples:
            document.add_paragraph(json.dumps(sample, indent=2, default=str))
    else:
        document.add_paragraph("No sample collection diagnostics were available.")


def _safe_int(value: Any, default: int = 0) -> int:
    """Convert values to int without leaking parsing errors."""

    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _slugify(value: str) -> str:
    """Return a conservative filename slug."""

    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value).strip("-").lower()
    return slug or "chart"
