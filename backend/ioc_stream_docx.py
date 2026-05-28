"""DOCX export helpers for GTI Recent IoC Stream Sample reports."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from backend.top_ranking_docx import (  # noqa: PLC0415
    _build_bar_chart_xml as _bar_xml,
    _build_pie_chart_xml as _pie_xml,
    _build_donut_chart_xml as _donut_xml,
    _build_line_chart_xml as _line_xml,
    _build_area_chart_xml as _area_xml,
    _build_stacked_bar_chart_xml as _stacked_xml,
    _insert_native_chart,
    _safe_int,
)


def generate_ioc_stream_docx(report_data: dict[str, Any], output_path: str) -> str:
    """Render a Recent IoC Stream Sample report as a readable DOCX document."""

    document = Document()
    core = document.core_properties
    core.title = "GTI Recent IoC Stream Sample Report"
    core.subject = "Client-friendly summary of GTI IoC Stream notifications"

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    summary = _as_dict(report_data.get("summary"))
    document.add_heading("GTI Recent IoC Stream Sample Report", level=0)
    document.add_paragraph(f"Generated at: {summary.get('generated_at') or 'Unknown'}")

    document.add_heading("Executive Summary", level=1)
    for sentence in _as_list(report_data.get("business_summary")):
        document.add_paragraph(str(sentence), style="List Bullet")

    document.add_heading("Key Metrics", level=1)
    technical_details = _as_dict(report_data.get("technical_details"))
    enrichment = _as_dict(technical_details.get("enrichment"))
    collection = _as_dict(report_data.get("collection") or technical_details.get("collection"))
    metrics = [
        ("Collection mode", collection.get("collection_mode_label", collection.get("collection_mode", "recent_pages"))),
        ("Time window", _as_dict(collection.get("time_window")).get("label", "n/a")),
        ("Requested pages", collection.get("requested_pages", "Unknown")),
        ("Max pages", collection.get("max_pages", collection.get("requested_pages", "Unknown"))),
        ("Pages fetched", collection.get("pages_fetched", 0)),
        ("Page size", collection.get("page_size", 40)),
        ("Raw IoCs fetched", collection.get("raw_ioc_count", summary.get("raw_ioc_count", 0))),
        ("IoCs returned by GTI date filter", collection.get("raw_iocs_returned", collection.get("server_side_date_filter_returned_count", collection.get("raw_ioc_count", summary.get("raw_ioc_count", 0))))),
        ("Unique IoCs after deduplication", collection.get("unique_ioc_count", summary.get("total_iocs", 0))),
        ("Duplicates removed", collection.get("duplicates_removed", 0)),
        ("Stopped reason", collection.get("stopped_reason", "unknown")),
        ("Coverage status", collection.get("coverage_status", "n/a")),
        ("Server date filter string", collection.get("server_side_date_filter_string") or "n/a"),
        ("Server date filter returned count", collection.get("server_side_date_filter_returned_count", collection.get("server_side_date_filter_item_count", 0))),
        ("GTI date filter sent", collection.get("server_side_date_filter_string") or "n/a"),
        ("Recommendation", collection.get("recommendation") or "n/a"),
        ("Earliest fetched timestamp", collection.get("earliest_fetched_timestamp") or "n/a"),
        ("Latest fetched timestamp", collection.get("latest_fetched_timestamp") or "n/a"),
        ("Earliest parsed timestamp (diagnostic only)", collection.get("earliest_fetched_timestamp") or "n/a"),
        ("Latest parsed timestamp (diagnostic only)", collection.get("latest_fetched_timestamp") or "n/a"),
        ("Items with stream timestamp", collection.get("items_with_stream_timestamp", 0)),
        ("Items without stream timestamp", collection.get("items_without_stream_timestamp", collection.get("items_without_stream_timestamp_count", 0))),
        ("Stream timestamp fields seen", ", ".join(str(field) for field in _as_list(collection.get("stream_timestamp_fields_seen"))) or "n/a"),
        ("Object metadata timestamp fields seen", ", ".join(str(field) for field in _as_list(collection.get("object_metadata_timestamp_fields_seen"))) or "n/a"),
        ("Timestamp fields seen", ", ".join(str(field) for field in _as_list(collection.get("timestamp_fields_seen"))) or "n/a"),
        ("Stop timestamp field", collection.get("stop_timestamp_field") or "n/a"),
        ("Oldest stream event timestamp", collection.get("oldest_stream_event_timestamp") or "n/a"),
        ("Oldest object metadata timestamp", collection.get("oldest_object_metadata_timestamp") or "n/a"),
        ("Ignored old object metadata timestamps", collection.get("ignored_object_metadata_old_timestamp_count", 0)),
        ("Items without stream timestamp", collection.get("items_without_stream_timestamp_count", 0)),
        ("Total IoCs", summary.get("total_iocs", 0)),
        ("Total enriched", collection.get("total_enriched", enrichment.get("succeeded", 0))),
        ("High risk", summary.get("high_risk", 0)),
        ("Medium risk", summary.get("medium_risk", 0)),
        ("Low risk", summary.get("low_risk", 0)),
        ("Unknown risk", summary.get("unknown_risk", 0)),
        ("Main IoC type", summary.get("main_entity_type", "Unknown")),
        ("Main source type", summary.get("main_source_type", "Unknown")),
        ("Enrichment enabled", enrichment.get("enabled", False)),
        ("Enrichment attempted", enrichment.get("attempted", 0)),
        ("Enrichment succeeded", enrichment.get("succeeded", 0)),
        ("Enrichment skipped", enrichment.get("skipped", 0)),
        ("Enrichment errors", enrichment.get("errors", 0)),
        ("Enrichment skipped too-long URLs", enrichment.get("skipped_too_long_url", 0)),
        ("Enrichment requested", enrichment.get("requested_limit", 0)),
        ("Enrichment actual scope", enrichment.get("actual_limit", 0)),
    ]
    _add_key_value_table(document, metrics)
    document.add_paragraph(
        "Risk scoring requires enrichment. This may generate one API lookup per IoC. "
        "Unknown means enrichment failed or GTI returned no risk context."
    )
    warnings = _as_list(technical_details.get("warnings"))
    if warnings:
        document.add_heading("Collection Warnings", level=1)
        for warning in warnings:
            document.add_paragraph(str(warning), style="List Bullet")

    document.add_heading("Chart Data", level=1)
    charts = _as_dict(report_data.get("charts"))

    def _insert_chart_multi(title: str, rows: Any, primary: str) -> None:
        """Insert a primary chart plus complementary chart types for the same data."""
        chart_rows = _as_list(rows)
        labels = [str(_as_dict(r).get("label") or "Unknown") for r in chart_rows]
        values = [_safe_int(_as_dict(r).get("value", 0)) for r in chart_rows]
        document.add_heading(title, level=2)
        if not labels or not any(v > 0 for v in values):
            _add_label_value_table(document, chart_rows)
            return
        bar_h = max(2.0, min(4.0, len(labels) * 0.35 + 0.8))
        try:
            if primary == "pie":
                _insert_native_chart(document, _pie_xml(labels, values, title), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _donut_xml(labels, values, title + " — Donut"), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _bar_xml(labels, values, title + " — Bars", horizontal=True), width_inches=5.5, height_inches=bar_h)
            else:  # bar
                _insert_native_chart(document, _bar_xml(labels, values, title, horizontal=True), width_inches=5.5, height_inches=bar_h)
                _insert_native_chart(document, _bar_xml(labels, values, title + " — Columns", horizontal=False), width_inches=5.5, height_inches=3.0)
                _insert_native_chart(document, _pie_xml(labels, values, title + " — Pie"), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _donut_xml(labels, values, title + " — Donut"), width_inches=5.0, height_inches=3.0)
        except Exception:
            _add_label_value_table(document, chart_rows)

    _insert_chart_multi("IoCs by entity type", charts.get("by_entity_type"), "bar")
    _insert_chart_multi("IoCs by risk", charts.get("by_risk"), "pie")
    _insert_chart_multi("IoCs by source type", charts.get("by_source_type"), "bar")
    _insert_chart_multi("Recommended actions", charts.get("by_recommended_action"), "bar")

    # Enrichment Outcome pie/donut
    enrich_succeeded = _safe_int(_as_dict(_as_dict(report_data.get("technical_details")).get("enrichment")).get("succeeded"))
    enrich_skipped = _safe_int(_as_dict(_as_dict(report_data.get("technical_details")).get("enrichment")).get("skipped"))
    enrich_errors = _safe_int(_as_dict(_as_dict(report_data.get("technical_details")).get("enrichment")).get("errors"))
    enrich_labels = ["Succeeded", "Skipped", "Errors"]
    enrich_values = [enrich_succeeded, enrich_skipped, enrich_errors]
    if any(v > 0 for v in enrich_values):
        document.add_heading("Enrichment Outcome", level=2)
        try:
            _insert_native_chart(document, _pie_xml(enrich_labels, enrich_values, "Enrichment Outcome"), width_inches=5.0, height_inches=3.0)
            _insert_native_chart(document, _donut_xml(enrich_labels, enrich_values, "Enrichment Outcome — Donut"), width_inches=5.0, height_inches=3.0)
        except Exception:
            pass

    analytics = _as_dict(report_data.get("analytics"))
    document.add_heading("Analyst Cross-Analysis", level=1)
    document.add_paragraph(
        "Computed only from successfully enriched indicators. Missing fields are left as n/a."
    )
    _add_key_value_table(
        document,
        [
            ("Analytics source", analytics.get("source", "enriched_indicators_only")),
            ("Enriched indicators analyzed", analytics.get("enriched_indicator_count", 0)),
        ],
    )
    document.add_heading("Business Insights", level=2)
    for sentence in _as_list(analytics.get("business_insights")):
        document.add_paragraph(str(sentence), style="List Bullet")
    document.add_heading("Highest Risk by IoC Type", level=2)
    _add_ioc_type_risk_table(document, _as_list(analytics.get("highest_risk_by_ioc_type")))
    document.add_heading("Top 10 Most Dangerous Indicators", level=2)
    _add_dangerous_indicators_table(
        document,
        _as_list(analytics.get("top_dangerous_indicators")),
    )
    def _insert_dist_multi(title: str, rows: Any, primary: str) -> None:
        dist_rows = _as_list(rows)
        labels = [str(_as_dict(r).get("label") or "Unknown") for r in dist_rows]
        values = [_safe_int(_as_dict(r).get("count", _as_dict(r).get("value", 0))) for r in dist_rows]
        document.add_heading(title, level=2)
        if not labels or not any(v > 0 for v in values):
            _add_distribution_table(document, dist_rows)
            return
        bar_h = max(2.0, min(4.0, len(labels) * 0.35 + 0.8))
        try:
            if primary == "pie":
                _insert_native_chart(document, _pie_xml(labels, values, title), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _donut_xml(labels, values, title + " — Donut"), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _bar_xml(labels, values, title + " — Bars", horizontal=True), width_inches=5.5, height_inches=bar_h)
            else:
                _insert_native_chart(document, _bar_xml(labels, values, title, horizontal=True), width_inches=5.5, height_inches=bar_h)
                _insert_native_chart(document, _bar_xml(labels, values, title + " — Columns", horizontal=False), width_inches=5.5, height_inches=3.0)
                _insert_native_chart(document, _pie_xml(labels, values, title + " — Pie"), width_inches=5.0, height_inches=3.0)
                _insert_native_chart(document, _donut_xml(labels, values, title + " — Donut"), width_inches=5.0, height_inches=3.0)
        except Exception:
            _add_distribution_table(document, dist_rows)

    _insert_dist_multi("Risk Distribution", analytics.get("risk_distribution"), "pie")
    _insert_dist_multi("IoC Type Distribution", analytics.get("ioc_type_distribution"), "bar")
    _insert_dist_multi("Recommended Action Distribution", analytics.get("recommended_action_distribution"), "bar")

    # Stacked bar: Malicious vs Suspicious per IoC type
    type_risk_rows = _as_list(analytics.get("highest_risk_by_ioc_type"))
    if type_risk_rows:
        ioc_types = [str(_as_dict(r).get("ioc_type") or "others") for r in type_risk_rows]
        malicious_vals = [_safe_int(_as_dict(r).get("malicious_indicator_count", 0)) for r in type_risk_rows]
        suspicious_vals = [_safe_int(_as_dict(r).get("suspicious_indicator_count", 0)) for r in type_risk_rows]
        avg_scores = [_safe_int(_as_dict(r).get("average_risk_score", 0)) for r in type_risk_rows]
        malicious_pct = [_safe_int(_as_dict(r).get("malicious_percentage", 0)) for r in type_risk_rows]

        if any(v > 0 for v in malicious_vals + suspicious_vals):
            document.add_heading("Malicious vs Suspicious per IoC Type", level=2)
            try:
                _insert_native_chart(
                    document,
                    _stacked_xml(
                        categories=ioc_types,
                        series_names=["Malicious", "Suspicious"],
                        series_values_list=[malicious_vals, suspicious_vals],
                        title="Malicious vs Suspicious per IoC Type",
                        percent_stacked=False,
                        horizontal=True,
                        series_colors=["DC3545", "FD7E14"],
                    ),
                    width_inches=6.0,
                    height_inches=max(2.5, len(ioc_types) * 0.4 + 0.8),
                )
                # 100% stacked variant
                _insert_native_chart(
                    document,
                    _stacked_xml(
                        categories=ioc_types,
                        series_names=["Malicious", "Suspicious"],
                        series_values_list=[malicious_vals, suspicious_vals],
                        title="Malicious vs Suspicious per IoC Type — 100%",
                        percent_stacked=True,
                        horizontal=True,
                        series_colors=["DC3545", "FD7E14"],
                    ),
                    width_inches=6.0,
                    height_inches=max(2.5, len(ioc_types) * 0.4 + 0.8),
                )
            except Exception:
                pass

        if any(v > 0 for v in malicious_pct):
            document.add_heading("Malicious % per IoC Type", level=2)
            try:
                _insert_native_chart(
                    document,
                    _bar_xml(ioc_types, malicious_pct, "Malicious % per IoC Type", horizontal=True),
                    width_inches=5.5,
                    height_inches=max(2.0, len(ioc_types) * 0.35 + 0.8),
                )
                _insert_native_chart(
                    document,
                    _bar_xml(ioc_types, malicious_pct, "Malicious % per IoC Type — Columns", horizontal=False),
                    width_inches=5.5,
                    height_inches=3.0,
                )
            except Exception:
                pass

        if any(v > 0 for v in avg_scores):
            document.add_heading("Average GTI Score per IoC Type", level=2)
            try:
                _insert_native_chart(
                    document,
                    _bar_xml(ioc_types, avg_scores, "Avg GTI Score per IoC Type", horizontal=True),
                    width_inches=5.5,
                    height_inches=max(2.0, len(ioc_types) * 0.35 + 0.8),
                )
                _insert_native_chart(
                    document,
                    _bar_xml(ioc_types, avg_scores, "Avg GTI Score per IoC Type — Columns", horizontal=False),
                    width_inches=5.5,
                    height_inches=3.0,
                )
            except Exception:
                pass

    document.add_heading("Top Indicators Requiring Attention", level=1)
    _add_top_indicators_table(document, _as_list(report_data.get("top_indicators"))[:10])

    document.add_heading("Business Interpretation", level=1)
    for sentence in _as_list(report_data.get("business_summary")):
        document.add_paragraph(str(sentence), style="List Bullet")

    document.add_heading("Definitions", level=1)
    _add_definitions_table(document, _as_list(report_data.get("definitions")))

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return str(output)


def _add_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)


def _add_label_value_table(document: Document, rows: list[Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Label"
    table.rows[0].cells[1].text = "Value"
    if not rows:
        cells = table.add_row().cells
        cells[0].text = "No data"
        cells[1].text = "0"
        return
    for row in rows:
        row_dict = _as_dict(row)
        cells = table.add_row().cells
        cells[0].text = str(row_dict.get("label", "Unknown"))
        cells[1].text = str(row_dict.get("value", 0))


def _add_top_indicators_table(document: Document, indicators: list[Any]) -> None:
    headers = (
        "Indicator",
        "Type",
        "Risk",
        "Malicious",
        "Suspicious",
        "Reputation",
        "Source",
        "Recommended action",
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    if not indicators:
        cells = table.add_row().cells
        cells[0].text = "No IoC Stream notifications were returned."
        return
    for indicator in indicators:
        item = _as_dict(indicator)
        cells = table.add_row().cells
        cells[0].text = str(item.get("display_value") or item.get("value") or "Unknown")
        cells[1].text = str(item.get("entity_type") or "Unknown")
        cells[2].text = str(item.get("severity") or "Unknown")
        cells[3].text = _format_optional(item.get("malicious"))
        cells[4].text = _format_optional(item.get("suspicious"))
        cells[5].text = _format_optional(item.get("reputation"))
        cells[6].text = str(item.get("source_name") or item.get("source_type") or "Unknown")
        cells[7].text = str(item.get("recommended_action") or "Manual review")


def _add_ioc_type_risk_table(document: Document, rows: list[Any]) -> None:
    headers = (
        "IoC type",
        "Total",
        "Avg GTI score",
        "Malicious IoCs",
        "Suspicious IoCs",
        "Malicious %",
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    if not rows:
        table.add_row().cells[0].text = "No enriched IoCs available."
        return
    for row in rows:
        item = _as_dict(row)
        cells = table.add_row().cells
        cells[0].text = str(item.get("ioc_type") or "others")
        cells[1].text = str(item.get("total_count", 0))
        cells[2].text = _format_optional(item.get("average_risk_score"))
        cells[3].text = str(item.get("malicious_indicator_count", 0))
        cells[4].text = str(item.get("suspicious_indicator_count", 0))
        cells[5].text = f"{item.get('malicious_percentage', 0)}%"


def _add_dangerous_indicators_table(document: Document, indicators: list[Any]) -> None:
    headers = (
        "Indicator",
        "Type",
        "Malicious",
        "Suspicious",
        "Reputation",
        "Recommended action",
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    if not indicators:
        table.add_row().cells[0].text = "No enriched IoCs available."
        return
    for indicator in indicators:
        item = _as_dict(indicator)
        cells = table.add_row().cells
        cells[0].text = str(item.get("display_value") or item.get("indicator") or "Unknown")
        cells[1].text = str(item.get("type") or "Unknown")
        cells[2].text = str(item.get("malicious", 0))
        cells[3].text = str(item.get("suspicious", 0))
        cells[4].text = _format_optional(item.get("reputation"))
        cells[5].text = str(item.get("recommended_action") or "Manual review")


def _add_distribution_table(document: Document, rows: list[Any]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Bucket"
    table.rows[0].cells[1].text = "Count"
    table.rows[0].cells[2].text = "Percentage"
    if not rows:
        cells = table.add_row().cells
        cells[0].text = "No data"
        cells[1].text = "0"
        cells[2].text = "0%"
        return
    for row in rows:
        item = _as_dict(row)
        cells = table.add_row().cells
        cells[0].text = str(item.get("label") or "Unknown")
        cells[1].text = str(item.get("count", item.get("value", 0)))
        cells[2].text = f"{item.get('percentage', 0)}%"


def _add_definitions_table(document: Document, definitions: list[Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    for definition in definitions:
        item = _as_dict(definition)
        cells = table.add_row().cells
        cells[0].text = str(item.get("term") or "")
        cells[1].text = str(item.get("definition") or "")


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _format_optional(value: Any) -> str:
    return "n/a" if value is None else str(value)


def _format_actual_collection_window(collection: dict[str, Any]) -> str:
    earliest = collection.get("earliest_timestamp") or "n/a"
    latest = collection.get("latest_timestamp") or "n/a"
    return f"{earliest} -> {latest}"
