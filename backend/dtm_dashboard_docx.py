"""DOCX export helpers for GTI DTM Monitor & Alert Dashboard reports."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from backend.top_ranking_docx import (  # noqa: PLC0415
    _build_bar_chart_xml as _bar_xml,
    _build_pie_chart_xml as _pie_xml,
    _build_donut_chart_xml as _donut_xml,
    _build_line_chart_xml as _line_xml,
    _build_area_chart_xml as _area_xml,
    _build_stacked_bar_chart_xml as _stacked_xml,
    _insert_native_chart,
    _safe_int,
    _shade_cell,
)

# Severity → RGB for header cells
_SEVERITY_COLORS: dict[str, tuple[int, int, int]] = {
    "high": (220, 53, 69),
    "medium": (253, 126, 20),
    "low": (40, 167, 69),
}


def generate_dtm_dashboard_docx(
    dashboard_result: dict[str, Any],
    output_path: str,
    max_chart_items: int = 10,
) -> str:
    """Render a DTM Monitor & Alert Dashboard as a Word document."""

    document = Document()
    core = document.core_properties
    core.title = "GTI DTM Monitor & Alert Dashboard"
    core.subject = "Read-only summary of GTI DTM monitors and alerts"

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    period = _as_dict(dashboard_result.get("period"))
    summary = _as_dict(dashboard_result.get("summary"))
    quota = _as_dict(dashboard_result.get("quota"))
    charts = _as_dict(dashboard_result.get("charts"))
    monitors = _as_list(dashboard_result.get("monitors"))
    warnings = _as_list(dashboard_result.get("warnings"))
    limits = _as_dict(dashboard_result.get("limits"))

    # Build per-monitor chart rows directly from the full monitors list so that
    # max_chart_items is not limited by the [:10] hardcoded in the dashboard route.
    _monitors_by_alert = sorted(
        [m for m in monitors if _safe_int(_as_dict(m).get("alert_count")) > 0],
        key=lambda r: (-_safe_int(_as_dict(r).get("alert_count")), str(_as_dict(r).get("name", "")).casefold()),
    )
    _monitors_by_risk = sorted(
        [m for m in monitors if _safe_int(_as_dict(m).get("risk_score")) > 0],
        key=lambda r: (-_safe_int(_as_dict(r).get("risk_score")), str(_as_dict(r).get("name", "")).casefold()),
    )
    _monitors_by_noise = sorted(
        [m for m in monitors if _safe_int(_as_dict(m).get("noise_score")) > 0],
        key=lambda r: (-_safe_int(_as_dict(r).get("noise_score")), str(_as_dict(r).get("name", "")).casefold()),
    )

    def _alert_count_rows(n: int) -> list[dict]:
        return [{"monitor_name": _as_dict(r).get("name") or "Unknown",
                 "alert_count": _safe_int(_as_dict(r).get("alert_count"))}
                for r in _monitors_by_alert[:n]]

    def _risk_score_rows(n: int) -> list[dict]:
        return [{"monitor_name": _as_dict(r).get("name") or "Unknown",
                 "risk_score": _safe_int(_as_dict(r).get("risk_score")),
                 "high": _safe_int(_as_dict(r).get("high")),
                 "medium": _safe_int(_as_dict(r).get("medium")),
                 "low": _safe_int(_as_dict(r).get("low"))}
                for r in _monitors_by_risk[:n]]

    def _noise_score_rows(n: int) -> list[dict]:
        return [{"monitor_name": _as_dict(r).get("name") or "Unknown",
                 "noise_score": _safe_int(_as_dict(r).get("noise_score")),
                 "risk_score": _safe_int(_as_dict(r).get("risk_score"))}
                for r in _monitors_by_noise[:n]]
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # ── Cover ───────────────────────────────────────────────────────────────
    document.add_heading("DTM Monitor & Alert Dashboard", level=0)
    document.add_paragraph(
        f"Period: {period.get('since', 'n/a')}  →  {period.get('until', 'n/a')}"
    )
    document.add_paragraph(f"Generated at: {generated_at}")
    max_alerts = _safe_int(limits.get("max_alerts"))
    if max_alerts:
        document.add_paragraph(
            f"Showing up to {max_alerts} alerts. Increase max_pages to fetch more."
        )

    # ── Warnings ────────────────────────────────────────────────────────────
    if warnings:
        document.add_heading("Notices", level=1)
        for warning in warnings:
            document.add_paragraph(str(warning), style="List Bullet")

    # ── KPI summary ─────────────────────────────────────────────────────────
    document.add_heading("Summary", level=1)
    kpi_rows: list[tuple[str, Any]] = [
        ("Total monitors", summary.get("total_monitors", 0)),
        ("Total alerts", summary.get("total_alerts", 0)),
        ("High severity alerts", summary.get("high_alerts", 0)),
        ("Medium severity alerts", summary.get("medium_alerts", 0)),
        ("Low severity alerts", summary.get("low_alerts", 0)),
        ("Monitors with alerts", summary.get("monitors_with_alerts", 0)),
        ("Inactive monitors", summary.get("monitors_without_alerts", 0)),
        ("Top risky monitor", summary.get("top_risky_monitor") or "none"),
        ("Top noisy monitor", summary.get("top_noisy_monitor") or "none"),
        ("Monitor quota used", f"{quota.get('used_percent', 0):.1f}%"),
        ("Monitor count / quota", f"{quota.get('monitor_count', 0)} / {quota.get('default_monitor_quota', 100)}"),
        ("Remaining quota estimate", quota.get("remaining_estimate", 0)),
    ]
    _add_kv_table(document, kpi_rows)

    # ── Charts ──────────────────────────────────────────────────────────────
    document.add_page_break()
    document.add_heading("Charts", level=1)

    def _cs(title: str, rows: Any, label_key: str, value_key: str,
            chart_type: str = "bar", width: float = 6.0, height: float | None = None) -> None:
        """Shorthand: call _add_chart_section with the current max_chart_items."""
        _add_chart_section(document, title, rows, label_key, value_key,
                           chart_type=chart_type, width=width, height=height,
                           max_items=max_chart_items)

    # ── Severity distribution: pie + donut + bar ─────────────────────────────
    _cs("Alerts by Severity", charts.get("alerts_by_severity", []),
        "severity", "count", chart_type="pie", width=4.5, height=3.0)
    _cs("Alerts by Severity — Donut", charts.get("alerts_by_severity", []),
        "severity", "count", chart_type="donut", width=4.5, height=3.0)
    _cs("Alerts by Severity — Bars", charts.get("alerts_by_severity", []),
        "severity", "count", chart_type="bar", width=5.5)

    # ── Monitors with vs without alerts (new) ────────────────────────────────
    monitors_with = _safe_int(summary.get("monitors_with_alerts"))
    monitors_without = _safe_int(summary.get("monitors_without_alerts"))
    if monitors_with + monitors_without > 0:
        _cs("Monitors with Alerts vs Inactive",
            [{"label": "With alerts", "count": monitors_with},
             {"label": "Inactive", "count": monitors_without}],
            "label", "count", chart_type="pie", width=4.5, height=3.0)
        _cs("Monitors with Alerts vs Inactive — Donut",
            [{"label": "With alerts", "count": monitors_with},
             {"label": "Inactive", "count": monitors_without}],
            "label", "count", chart_type="donut", width=4.5, height=3.0)

    # ── Quota usage (new) ────────────────────────────────────────────────────
    quota_used = _safe_int(quota.get("monitor_count"))
    quota_remaining = _safe_int(quota.get("remaining_estimate"))
    if quota_used + quota_remaining > 0:
        _cs("Monitor Quota Usage",
            [{"label": "Used", "count": quota_used},
             {"label": "Remaining", "count": quota_remaining}],
            "label", "count", chart_type="donut", width=4.5, height=3.0)

    # ── Top Monitors by Alert Count ──────────────────────────────────────────
    _cs("Top Monitors by Alert Count", _alert_count_rows(max_chart_items),
        "monitor_name", "alert_count", chart_type="bar", width=6.0)
    _cs("Top Monitors by Alert Count — Columns", _alert_count_rows(max_chart_items),
        "monitor_name", "alert_count", chart_type="timeline", width=6.0, height=3.0)

    # ── Top Monitors by Risk Score ───────────────────────────────────────────
    _cs("Top Monitors by Risk Score", _risk_score_rows(max_chart_items),
        "monitor_name", "risk_score", chart_type="bar", width=6.0)

    # ── High / Medium / Low breakdown per monitor (stacked bars) ─────────────
    risk_rows = _risk_score_rows(max_chart_items)
    if risk_rows:
        risk_cats = [str(_as_dict(r).get("monitor_name") or "Unknown") for r in risk_rows]
        high_vals = [_safe_int(_as_dict(r).get("high", 0)) for r in risk_rows]
        med_vals = [_safe_int(_as_dict(r).get("medium", 0)) for r in risk_rows]
        low_vals = [_safe_int(_as_dict(r).get("low", 0)) for r in risk_rows]
        if any(v > 0 for v in high_vals + med_vals + low_vals):
            document.add_heading("Alerts per Monitor — High / Medium / Low", level=2)
            try:
                _insert_native_chart(
                    document,
                    _stacked_xml(
                        categories=risk_cats,
                        series_names=["High", "Medium", "Low"],
                        series_values_list=[high_vals, med_vals, low_vals],
                        title="Alerts per Monitor — Stacked",
                        percent_stacked=False,
                        horizontal=True,
                        series_colors=["DC3545", "FD7E14", "28A745"],
                    ),
                    width_inches=6.0,
                    height_inches=max(2.5, len(risk_cats) * 0.4 + 0.8),
                )
                _insert_native_chart(
                    document,
                    _stacked_xml(
                        categories=risk_cats,
                        series_names=["High", "Medium", "Low"],
                        series_values_list=[high_vals, med_vals, low_vals],
                        title="Alerts per Monitor — 100% Stacked",
                        percent_stacked=True,
                        horizontal=True,
                        series_colors=["DC3545", "FD7E14", "28A745"],
                    ),
                    width_inches=6.0,
                    height_inches=max(2.5, len(risk_cats) * 0.4 + 0.8),
                )
            except Exception:
                pass

    # ── Alerts by Type ───────────────────────────────────────────────────────
    _cs("Alerts by Type", charts.get("alerts_by_type", []),
        "type", "count", chart_type="bar", width=6.0)
    _cs("Alerts by Type — Pie", charts.get("alerts_by_type", []),
        "type", "count", chart_type="pie", width=5.0, height=3.0)
    _cs("Alerts by Type — Donut", charts.get("alerts_by_type", []),
        "type", "count", chart_type="donut", width=5.0, height=3.0)

    # ── Alerts by Status ─────────────────────────────────────────────────────
    _cs("Alerts by Status", charts.get("alerts_by_status", []),
        "status", "count", chart_type="bar", width=6.0)
    _cs("Alerts by Status — Pie", charts.get("alerts_by_status", []),
        "status", "count", chart_type="pie", width=5.0, height=3.0)
    _cs("Alerts by Status — Donut", charts.get("alerts_by_status", []),
        "status", "count", chart_type="donut", width=5.0, height=3.0)

    # ── Timeline: vertical bar + line + area ─────────────────────────────────
    _cs("Alerts Timeline (daily) — Columns", charts.get("alerts_timeline", []),
        "date", "count", chart_type="timeline", width=6.0, height=3.0)
    _cs("Alerts Timeline (daily) — Line", charts.get("alerts_timeline", []),
        "date", "count", chart_type="line", width=6.0, height=3.0)
    _cs("Alerts Timeline (daily) — Area", charts.get("alerts_timeline", []),
        "date", "count", chart_type="area", width=6.0, height=3.0)

    # ── Noisy monitors ───────────────────────────────────────────────────────
    _cs("Noisy Monitors — Noise Score", _noise_score_rows(max_chart_items),
        "monitor_name", "noise_score", chart_type="bar", width=6.0)

    # ── Risk Score vs Noise Score per monitor (clustered bar) ────────────────
    noisy_rows = _noise_score_rows(max_chart_items)
    if noisy_rows:
        noisy_cats = [str(_as_dict(r).get("monitor_name") or "Unknown") for r in noisy_rows]
        noisy_risk = [_safe_int(_as_dict(r).get("risk_score", 0)) for r in noisy_rows]
        noisy_noise = [_safe_int(_as_dict(r).get("noise_score", 0)) for r in noisy_rows]
        if any(v > 0 for v in noisy_risk + noisy_noise):
            document.add_heading("Risk Score vs Noise Score per Monitor", level=2)
            try:
                _insert_native_chart(
                    document,
                    _stacked_xml(
                        categories=noisy_cats,
                        series_names=["Risk Score", "Noise Score"],
                        series_values_list=[noisy_risk, noisy_noise],
                        title="Risk Score vs Noise Score",
                        percent_stacked=False,
                        horizontal=True,
                        series_colors=["0D7F7A", "6C757D"],
                    ),
                    width_inches=6.0,
                    height_inches=max(2.5, len(noisy_cats) * 0.4 + 0.8),
                )
            except Exception:
                pass

    # ── Monitor table ────────────────────────────────────────────────────────
    document.add_page_break()
    document.add_heading("Monitor Details", level=1)
    _add_monitor_table(document, monitors)

    # ── Inactive monitors ───────────────────────────────────────────────────
    inactive = _as_list(charts.get("inactive_monitors"))
    if inactive:
        document.add_heading("Inactive Monitors", level=1)
        _add_inactive_table(document, inactive)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return str(output)


# ── Chart helpers ─────────────────────────────────────────────────────────────

def _add_chart_section(
    document: Document,
    title: str,
    rows: Any,
    label_key: str,
    value_key: str,
    chart_type: str = "bar",
    width: float = 6.0,
    height: float | None = None,
    max_items: int | None = None,
) -> None:
    document.add_heading(title, level=2)
    data = _as_list(rows)
    if max_items is not None and chart_type not in ("pie", "donut", "timeline", "line", "area"):
        data = data[:max_items]
    labels = [str(_as_dict(r).get(label_key) or "Unknown") for r in data]
    values = [_safe_int(_as_dict(r).get(value_key, 0)) for r in data]

    if not labels or not any(v > 0 for v in values):
        document.add_paragraph("No data available for this period.")
        return

    try:
        auto_height = height or max(2.0, min(4.5, len(labels) * 0.35 + 0.8))
        if chart_type == "pie":
            xml = _pie_xml(labels, values, title)
            _insert_native_chart(document, xml, width_inches=width, height_inches=height or 3.0)
        elif chart_type == "donut":
            xml = _donut_xml(labels, values, title)
            _insert_native_chart(document, xml, width_inches=width, height_inches=height or 3.0)
        elif chart_type == "timeline":
            xml = _bar_xml(labels, values, title, horizontal=False)
            _insert_native_chart(document, xml, width_inches=width, height_inches=height or 3.0)
        elif chart_type == "line":
            xml = _line_xml(labels, values, title)
            _insert_native_chart(document, xml, width_inches=width, height_inches=height or 3.0)
        elif chart_type == "area":
            xml = _area_xml(labels, values, title)
            _insert_native_chart(document, xml, width_inches=width, height_inches=height or 3.0)
        else:
            xml = _bar_xml(labels, values, title, horizontal=True)
            _insert_native_chart(document, xml, width_inches=width, height_inches=auto_height)
    except Exception:
        # Fallback to a plain table if chart embedding fails.
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = label_key.replace("_", " ").title()
        table.rows[0].cells[1].text = value_key.replace("_", " ").title()
        for label, value in zip(labels, values):
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)


# ── Table helpers ─────────────────────────────────────────────────────────────

def _add_kv_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)


def _add_monitor_table(document: Document, monitors: list[Any]) -> None:
    if not monitors:
        document.add_paragraph("No monitor data available.")
        return

    headers = ("Monitor", "Alerts", "Risk", "Noise", "High", "Medium", "Low", "Last Alert")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    max_risk = max((_safe_int(_as_dict(m).get("risk_score")) for m in monitors), default=1) or 1
    for monitor in monitors:
        item = _as_dict(monitor)
        cells = table.add_row().cells
        cells[0].text = str(item.get("name") or "Unknown")
        cells[1].text = str(_safe_int(item.get("alert_count")))
        risk = _safe_int(item.get("risk_score"))
        cells[2].text = str(risk)
        cells[3].text = str(_safe_int(item.get("noise_score")))
        cells[4].text = str(_safe_int(item.get("high")))
        cells[5].text = str(_safe_int(item.get("medium")))
        cells[6].text = str(_safe_int(item.get("low")))
        cells[7].text = str(item.get("last_alert_date") or "none")
        # Heat-map shade on risk column
        intensity = min(1.0, risk / max_risk)
        r = round(255 - (255 - 220) * intensity * 0.7)
        g = round(255 - (255 - 53) * intensity * 0.7)
        b = round(255 - (255 - 69) * intensity * 0.7)
        fill = f"{r:02X}{g:02X}{b:02X}"
        if risk > 0:
            _shade_cell(cells[2], fill)


def _add_inactive_table(document: Document, monitors: list[Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Monitor"
    table.rows[0].cells[1].text = "Last Alert"
    for item in monitors[:30]:
        row = _as_dict(item)
        cells = table.add_row().cells
        cells[0].text = str(row.get("monitor_name") or "Unknown")
        cells[1].text = str(row.get("last_alert_date") or "none")


# ── Generic helpers ───────────────────────────────────────────────────────────

def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []
